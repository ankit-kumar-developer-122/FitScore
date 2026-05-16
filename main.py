from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
import os, sqlite3, hashlib, uuid, json, re
from datetime import datetime
from html.parser import HTMLParser
from urllib import request as urlrequest
from urllib.error import URLError, HTTPError
from urllib.parse import urljoin

from pypdf import PdfReader
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

app = FastAPI(title="FitScore API")

allowed_origins_env = os.getenv("ALLOWED_ORIGINS", "*")
allowed_origins = ["*"] if allowed_origins_env.strip() == "*" else [
    origin.strip() for origin in allowed_origins_env.split(",") if origin.strip()
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class CachedStaticFiles(StaticFiles):
    async def get_response(self, path, scope):
        response = await super().get_response(path, scope)
        if response.status_code >= 400:
            return response

        if path.endswith((".css", ".js")):
            response.headers["Cache-Control"] = "no-cache"
        elif path.endswith(".html") or path == ".":
            response.headers["Cache-Control"] = "no-cache"
        else:
            response.headers["Cache-Control"] = "public, max-age=3600"
        return response

# Directories
os.makedirs("public/css", exist_ok=True)
os.makedirs("public/js", exist_ok=True)
os.makedirs("uploads", exist_ok=True)
os.makedirs("uploads/generated", exist_ok=True)

# ── Database Setup ───────────────────────────────────────────────────
DB_PATH = "fitscore.db"
N8N_WEBHOOK_URL = os.getenv("N8N_WEBHOOK_URL", "").strip()
N8N_WEBHOOK_SECRET = os.getenv("N8N_WEBHOOK_SECRET", "").strip()
MATCH_THRESHOLD = int(os.getenv("JOB_MATCH_THRESHOLD", "60"))

TOP_COMPANY_SEEDS = [
    ("Google", "google.com", "https://www.google.com/about/careers/applications/"),
    ("Microsoft", "microsoft.com", "https://jobs.careers.microsoft.com/"),
    ("Amazon", "amazon.jobs", "https://www.amazon.jobs/"),
    ("Apple", "apple.com", "https://jobs.apple.com/"),
    ("Meta", "metacareers.com", "https://www.metacareers.com/jobs/"),
    ("Netflix", "netflix.com", "https://jobs.netflix.com/"),
    ("NVIDIA", "nvidia.com", "https://www.nvidia.com/en-us/about-nvidia/careers/"),
    ("OpenAI", "openai.com", "https://openai.com/careers/"),
    ("Anthropic", "anthropic.com", "https://www.anthropic.com/careers"),
    ("DeepMind", "deepmind.google", "https://deepmind.google/about/careers/"),
    ("Stripe", "stripe.com", "https://stripe.com/jobs"),
    ("Databricks", "databricks.com", "https://www.databricks.com/company/careers"),
    ("Snowflake", "snowflake.com", "https://careers.snowflake.com/"),
    ("MongoDB", "mongodb.com", "https://www.mongodb.com/careers"),
    ("Cloudflare", "cloudflare.com", "https://www.cloudflare.com/careers/jobs/"),
    ("Salesforce", "salesforce.com", "https://careers.salesforce.com/"),
    ("Adobe", "adobe.com", "https://careers.adobe.com/"),
    ("ServiceNow", "servicenow.com", "https://careers.servicenow.com/"),
    ("Atlassian", "atlassian.com", "https://www.atlassian.com/company/careers"),
    ("Uber", "uber.com", "https://www.uber.com/us/en/careers/"),
    ("Airbnb", "airbnb.com", "https://careers.airbnb.com/"),
    ("LinkedIn", "linkedin.com", "https://careers.linkedin.com/"),
    ("GitHub", "github.com", "https://github.com/about/careers"),
    ("GitLab", "gitlab.com", "https://about.gitlab.com/jobs/"),
    ("Figma", "figma.com", "https://www.figma.com/careers/"),
    ("Canva", "canva.com", "https://www.canva.com/careers/"),
    ("Shopify", "shopify.com", "https://www.shopify.com/careers"),
    ("PayPal", "paypal.com", "https://careers.pypl.com/"),
    ("Block", "block.xyz", "https://block.xyz/careers"),
    ("Coinbase", "coinbase.com", "https://www.coinbase.com/careers"),
    ("Robinhood", "robinhood.com", "https://careers.robinhood.com/"),
    ("Rippling", "rippling.com", "https://www.rippling.com/careers"),
    ("Brex", "brex.com", "https://www.brex.com/careers"),
    ("Ramp", "ramp.com", "https://ramp.com/careers"),
    ("Plaid", "plaid.com", "https://plaid.com/careers/"),
    ("Twilio", "twilio.com", "https://www.twilio.com/en-us/company/jobs"),
    ("Okta", "okta.com", "https://www.okta.com/company/careers/"),
    ("Zoom", "zoom.us", "https://careers.zoom.us/"),
    ("Slack", "salesforce.com", "https://slack.com/careers"),
    ("Oracle", "oracle.com", "https://www.oracle.com/careers/"),
    ("IBM", "ibm.com", "https://www.ibm.com/careers"),
    ("Intel", "intel.com", "https://jobs.intel.com/"),
    ("AMD", "amd.com", "https://www.amd.com/en/corporate/careers.html"),
    ("Qualcomm", "qualcomm.com", "https://www.qualcomm.com/company/careers"),
    ("Tesla", "tesla.com", "https://www.tesla.com/careers"),
    ("SpaceX", "spacex.com", "https://www.spacex.com/careers/"),
    ("Palantir", "palantir.com", "https://www.palantir.com/careers/"),
    ("DoorDash", "doordash.com", "https://careers.doordash.com/"),
    ("Instacart", "instacart.com", "https://www.instacart.com/company/careers"),
    ("Pinterest", "pinterest.com", "https://www.pinterestcareers.com/"),
    ("Reddit", "redditinc.com", "https://www.redditinc.com/careers"),
    ("Discord", "discord.com", "https://discord.com/careers"),
    ("Dropbox", "dropbox.com", "https://jobs.dropbox.com/"),
    ("Box", "box.com", "https://careers.box.com/"),
    ("Asana", "asana.com", "https://asana.com/jobs"),
    ("Notion", "notion.so", "https://www.notion.so/careers"),
    ("Airtable", "airtable.com", "https://www.airtable.com/careers"),
    ("Miro", "miro.com", "https://miro.com/careers/"),
    ("Grammarly", "grammarly.com", "https://www.grammarly.com/jobs"),
    ("Duolingo", "duolingo.com", "https://careers.duolingo.com/"),
    ("Coursera", "coursera.org", "https://www.coursera.org/careers"),
    ("Khan Academy", "khanacademy.org", "https://www.khanacademy.org/careers"),
    ("Workday", "workday.com", "https://workday.wd5.myworkdayjobs.com/Workday"),
    ("SAP", "sap.com", "https://jobs.sap.com/"),
    ("Cisco", "cisco.com", "https://jobs.cisco.com/"),
    ("VMware", "broadcom.com", "https://www.broadcom.com/company/careers"),
    ("Palo Alto Networks", "paloaltonetworks.com", "https://jobs.paloaltonetworks.com/"),
    ("CrowdStrike", "crowdstrike.com", "https://www.crowdstrike.com/careers/"),
    ("Zscaler", "zscaler.com", "https://www.zscaler.com/careers"),
    ("Datadog", "datadoghq.com", "https://careers.datadoghq.com/"),
    ("New Relic", "newrelic.com", "https://newrelic.com/careers"),
    ("Elastic", "elastic.co", "https://www.elastic.co/careers"),
    ("Confluent", "confluent.io", "https://www.confluent.io/careers/"),
    ("HashiCorp", "hashicorp.com", "https://www.hashicorp.com/careers"),
    ("DigitalOcean", "digitalocean.com", "https://www.digitalocean.com/careers"),
    ("Vercel", "vercel.com", "https://vercel.com/careers"),
    ("Netlify", "netlify.com", "https://www.netlify.com/careers/"),
    ("Supabase", "supabase.com", "https://supabase.com/careers"),
    ("PlanetScale", "planetscale.com", "https://planetscale.com/careers"),
    ("Postman", "postman.com", "https://www.postman.com/company/careers/"),
    ("Hugging Face", "huggingface.co", "https://huggingface.co/jobs"),
    ("Scale AI", "scale.com", "https://scale.com/careers"),
    ("Perplexity AI", "perplexity.ai", "https://www.perplexity.ai/hub/careers"),
    ("Mistral AI", "mistral.ai", "https://mistral.ai/careers"),
    ("Cohere", "cohere.com", "https://cohere.com/careers"),
    ("Character.AI", "character.ai", "https://character.ai/careers"),
    ("xAI", "x.ai", "https://x.ai/careers"),
    ("TCS", "tcs.com", "https://www.tcs.com/careers"),
    ("Infosys", "infosys.com", "https://www.infosys.com/careers/"),
    ("Wipro", "wipro.com", "https://careers.wipro.com/"),
    ("HCLTech", "hcltech.com", "https://www.hcltech.com/careers"),
    ("Tech Mahindra", "techmahindra.com", "https://www.techmahindra.com/en-in/careers/"),
    ("Accenture", "accenture.com", "https://www.accenture.com/us-en/careers"),
    ("Deloitte", "deloitte.com", "https://www.deloitte.com/global/en/careers.html"),
    ("EY", "ey.com", "https://www.ey.com/en_gl/careers"),
    ("KPMG", "kpmg.com", "https://kpmg.com/xx/en/home/careers.html"),
    ("PwC", "pwc.com", "https://www.pwc.com/gx/en/careers.html"),
    ("JPMorgan Chase", "jpmorganchase.com", "https://careers.jpmorgan.com/"),
    ("Goldman Sachs", "goldmansachs.com", "https://www.goldmansachs.com/careers/"),
    ("Morgan Stanley", "morganstanley.com", "https://www.morganstanley.com/careers"),
    ("Visa", "visa.com", "https://usa.visa.com/careers.html"),
    ("Mastercard", "mastercard.com", "https://careers.mastercard.com/"),
    ("American Express", "americanexpress.com", "https://www.americanexpress.com/en-us/careers/"),
    ("Walmart Global Tech", "walmart.com", "https://tech.walmart.com/content/walmart-global-tech/en_us/careers.html"),
    ("Target", "target.com", "https://jobs.target.com/"),
    ("Intuit", "intuit.com", "https://www.intuit.com/careers/"),
    ("Expedia Group", "expediagroup.com", "https://careers.expediagroup.com/"),
    ("Booking.com", "booking.com", "https://careers.booking.com/"),
    ("MakeMyTrip", "makemytrip.com", "https://careers.makemytrip.com/"),
    ("Flipkart", "flipkartcareers.com", "https://www.flipkartcareers.com/"),
    ("Myntra", "myntra.com", "https://www.myntra.com/careers"),
    ("Meesho", "meesho.io", "https://www.meesho.io/careers"),
    ("Swiggy", "swiggy.com", "https://careers.swiggy.com/"),
    ("Zomato", "zomato.com", "https://www.zomato.com/careers"),
    ("Zepto", "zeptonow.com", "https://www.zeptonow.com/careers"),
    ("PhonePe", "phonepe.com", "https://www.phonepe.com/careers/"),
    ("Razorpay", "razorpay.com", "https://razorpay.com/jobs/"),
    ("Paytm", "paytm.com", "https://paytm.com/careers/"),
    ("CRED", "cred.club", "https://careers.cred.club/"),
    ("Groww", "groww.in", "https://groww.in/careers"),
    ("Zerodha", "zerodha.com", "https://zerodha.com/careers/"),
    ("Freshworks", "freshworks.com", "https://www.freshworks.com/company/careers/"),
    ("Zoho", "zoho.com", "https://www.zoho.com/careers/"),
    ("BrowserStack", "browserstack.com", "https://www.browserstack.com/careers"),
    ("Chargebee", "chargebee.com", "https://www.chargebee.com/careers/"),
    ("InMobi", "inmobi.com", "https://www.inmobi.com/company/careers/"),
    ("Dream11", "dream11.com", "https://www.dream11.com/careers"),
    ("ShareChat", "sharechat.com", "https://sharechat.com/careers"),
    ("Ola", "olaelectric.com", "https://olaelectric.com/careers"),
    ("OYO", "oyorooms.com", "https://www.oyorooms.com/careers/"),
    ("Udaan", "udaan.com", "https://udaan.com/careers"),
    ("Navi", "navi.com", "https://navi.com/careers"),
    ("Pine Labs", "pinelabs.com", "https://www.pinelabs.com/careers"),
    ("Urban Company", "urbancompany.com", "https://www.urbancompany.com/careers"),
]

JOB_PORTAL_SEEDS = [
    ("Wellfound", "wellfound.com", "https://wellfound.com/jobs"),
    ("Y Combinator Work at a Startup", "ycombinator.com", "https://www.ycombinator.com/jobs"),
    ("Remote OK", "remoteok.com", "https://remoteok.com/remote-dev-jobs"),
    ("Hacker News Who is Hiring", "news.ycombinator.com", "https://news.ycombinator.com/jobs"),
]

JOB_TITLE_KEYWORDS = [
    "software engineer", "frontend", "backend", "full stack", "fullstack",
    "data scientist", "data engineer", "machine learning", "ml engineer",
    "ai engineer", "devops", "sre", "cloud engineer", "security engineer",
    "android", "ios", "product designer", "product manager", "intern",
]

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def get_request_user(request: Request) -> sqlite3.Row:
    user_id = (request.headers.get("X-User-Id") or "").strip()
    if not user_id:
        raise HTTPException(401, "Authentication required")
    conn = get_db()
    user = conn.execute("SELECT * FROM users WHERE id=?", (user_id,)).fetchone()
    conn.close()
    if not user:
        raise HTTPException(401, "User not found")
    return user


def require_roles(request: Request, allowed_roles: set[str]) -> sqlite3.Row:
    user = get_request_user(request)
    if user["role"] not in allowed_roles:
        raise HTTPException(403, "You do not have permission to perform this action")
    return user

def init_db():
    conn = get_db()
    c = conn.cursor()
    c.executescript("""
    CREATE TABLE IF NOT EXISTS users (
        id TEXT PRIMARY KEY,
        name TEXT NOT NULL,
        email TEXT UNIQUE NOT NULL,
        password_hash TEXT NOT NULL,
        role TEXT DEFAULT 'candidate',
        fitscore INTEGER DEFAULT 0,
        ats_score INTEGER DEFAULT 0,
        skills TEXT DEFAULT '[]',
        auto_email INTEGER DEFAULT 0,
        created_at TEXT
    );
    CREATE TABLE IF NOT EXISTS jobs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT, company TEXT, salary TEXT,
        match_pct INTEGER DEFAULT 0, status TEXT DEFAULT 'active',
        skills TEXT DEFAULT '[]', experience TEXT,
        description TEXT, location TEXT DEFAULT '',
        package_text TEXT DEFAULT '', application_url TEXT DEFAULT '',
        source_name TEXT DEFAULT '', source_url TEXT DEFAULT '',
        created_at TEXT
    );
    CREATE TABLE IF NOT EXISTS company_watchlist (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT UNIQUE NOT NULL,
        domain TEXT DEFAULT '',
        careers_url TEXT DEFAULT '',
        package_tier TEXT DEFAULT 'high',
        source_type TEXT DEFAULT 'company',
        created_at TEXT
    );
    CREATE TABLE IF NOT EXISTS user_job_preferences (
        user_id TEXT PRIMARY KEY,
        company_ids TEXT DEFAULT '[]',
        min_salary_lpa INTEGER DEFAULT 0,
        roles TEXT DEFAULT '',
        locations TEXT DEFAULT '',
        updated_at TEXT,
        FOREIGN KEY(user_id) REFERENCES users(id)
    );
    CREATE TABLE IF NOT EXISTS applications (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id TEXT, job_id INTEGER,
        stage TEXT DEFAULT 'applied',
        created_at TEXT,
        FOREIGN KEY(user_id) REFERENCES users(id),
        FOREIGN KEY(job_id) REFERENCES jobs(id)
    );
    CREATE TABLE IF NOT EXISTS resumes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id TEXT, filename TEXT, filepath TEXT,
        version_tag TEXT DEFAULT 'Original',
        ats_score INTEGER DEFAULT 0,
        analysis_json TEXT DEFAULT '{}',
        uploaded_at TEXT,
        FOREIGN KEY(user_id) REFERENCES users(id)
    );
    CREATE TABLE IF NOT EXISTS workflows (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT, status TEXT DEFAULT 'active',
        runs INTEGER DEFAULT 0, success_rate REAL DEFAULT 100.0,
        created_at TEXT
    );
    CREATE TABLE IF NOT EXISTS notification_logs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id TEXT NOT NULL,
        job_id INTEGER NOT NULL,
        channel TEXT NOT NULL,
        status TEXT NOT NULL,
        payload_json TEXT DEFAULT '{}',
        created_at TEXT,
        UNIQUE(user_id, job_id, channel),
        FOREIGN KEY(user_id) REFERENCES users(id),
        FOREIGN KEY(job_id) REFERENCES jobs(id)
    );
    """)
    resume_columns = {row[1] for row in c.execute("PRAGMA table_info(resumes)").fetchall()}
    job_columns = {row[1] for row in c.execute("PRAGMA table_info(jobs)").fetchall()}
    if "ats_score" not in resume_columns:
        try:
            c.execute("ALTER TABLE resumes ADD COLUMN ats_score INTEGER DEFAULT 0")
        except sqlite3.OperationalError as exc:
            if "duplicate column name" not in str(exc).lower():
                raise
    if "analysis_json" not in resume_columns:
        try:
            c.execute("ALTER TABLE resumes ADD COLUMN analysis_json TEXT DEFAULT '{}'")
        except sqlite3.OperationalError as exc:
            if "duplicate column name" not in str(exc).lower():
                raise
    if "location" not in job_columns:
        try:
            c.execute("ALTER TABLE jobs ADD COLUMN location TEXT DEFAULT ''")
        except sqlite3.OperationalError as exc:
            if "duplicate column name" not in str(exc).lower():
                raise
    if "package_text" not in job_columns:
        try:
            c.execute("ALTER TABLE jobs ADD COLUMN package_text TEXT DEFAULT ''")
        except sqlite3.OperationalError as exc:
            if "duplicate column name" not in str(exc).lower():
                raise
    if "application_url" not in job_columns:
        try:
            c.execute("ALTER TABLE jobs ADD COLUMN application_url TEXT DEFAULT ''")
        except sqlite3.OperationalError as exc:
            if "duplicate column name" not in str(exc).lower():
                raise
    if "source_name" not in job_columns:
        try:
            c.execute("ALTER TABLE jobs ADD COLUMN source_name TEXT DEFAULT ''")
        except sqlite3.OperationalError as exc:
            if "duplicate column name" not in str(exc).lower():
                raise
    if "source_url" not in job_columns:
        try:
            c.execute("ALTER TABLE jobs ADD COLUMN source_url TEXT DEFAULT ''")
        except sqlite3.OperationalError as exc:
            if "duplicate column name" not in str(exc).lower():
                raise
    if c.execute("SELECT COUNT(*) FROM company_watchlist").fetchone()[0] == 0:
        now = datetime.utcnow().isoformat()
        company_rows = [
            (name, domain, careers_url, "premium" if index < 80 else "high", "company", now)
            for index, (name, domain, careers_url) in enumerate(TOP_COMPANY_SEEDS)
        ]
        portal_rows = [
            (name, domain, url, "varied", "portal", now)
            for name, domain, url in JOB_PORTAL_SEEDS
        ]
        c.executemany(
            "INSERT OR IGNORE INTO company_watchlist (name,domain,careers_url,package_tier,source_type,created_at) VALUES (?,?,?,?,?,?)",
            company_rows + portal_rows,
        )
    # Seed data if empty
    if c.execute("SELECT COUNT(*) FROM jobs").fetchone()[0] == 0:
        now = datetime.utcnow().isoformat()
        jobs = [
            ("Senior Frontend Engineer","Vercel","$150k",95,"active",'["JavaScript","React","CSS"]',"3+ yrs","Build next-gen web UIs","Remote","15-18 LPA","https://vercel.com/careers","Vercel","https://vercel.com/careers",now),
            ("Data Scientist","OpenAI","$180k",88,"active",'["Python","ML","Statistics"]',"2+ yrs","Work on cutting-edge AI models","San Francisco, CA","22-28 LPA","https://openai.com/careers","OpenAI","https://openai.com/careers",now),
            ("Backend Engineer","Stripe","$160k",82,"active",'["Python","Go","SQL"]',"4+ yrs","Design payment infrastructure","Bengaluru","18-24 LPA","https://stripe.com/jobs","Stripe","https://stripe.com/jobs",now),
            ("Product Designer","Figma","$140k",78,"active",'["Figma","UI/UX","Prototyping"]',"2+ yrs","Shape the future of design tools","London","14-18 LPA","https://www.figma.com/careers/","Figma","https://www.figma.com/careers/",now),
            ("ML Engineer","DeepMind","$200k",92,"active",'["Python","TensorFlow","Research"]',"3+ yrs","Push boundaries of AI research","Mountain View, CA","25-32 LPA","https://deepmind.google/about/careers/","DeepMind","https://deepmind.google/about/careers/",now),
            ("DevOps Lead","Netflix","$170k",75,"active",'["AWS","Kubernetes","Terraform"]',"5+ yrs","Scale global streaming infra","Remote","20-26 LPA","https://jobs.netflix.com/","Netflix","https://jobs.netflix.com/",now),
        ]
        c.executemany("INSERT INTO jobs (title,company,salary,match_pct,status,skills,experience,description,location,package_text,application_url,source_name,source_url,created_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)", jobs)
    if c.execute("SELECT COUNT(*) FROM workflows").fetchone()[0] == 0:
        now = datetime.utcnow().isoformat()
        wfs = [
            ("Daily Job Scraper","active",142,98.5,now),
            ("Resume Enhancer","paused",12,100.0,now),
            ("Auto Match Pipeline","active",89,96.2,now),
        ]
        c.executemany("INSERT INTO workflows (name,status,runs,success_rate,created_at) VALUES (?,?,?,?,?)", wfs)
    conn.commit()
    conn.close()

init_db()


ACTION_VERBS = {
    "built", "created", "developed", "designed", "implemented", "improved",
    "optimized", "led", "managed", "launched", "delivered", "automated",
    "analyzed", "architected", "collaborated", "scaled"
}

COMMON_KEYWORDS = {
    "python", "javascript", "react", "sql", "aws", "docker", "kubernetes",
    "ci/cd", "git", "api", "machine learning", "data", "fastapi", "testing"
}

REFERENCE_RESUME_SECTIONS = [
    "Education",
    "Competitive Programming & DSA",
    "Projects",
    "Technical Skills",
]

SECTION_PATTERNS = {
    "contact": r"(email|phone|linkedin|github)",
    "summary": r"(summary|profile|objective)",
    "experience": r"(experience|employment|work history)",
    "education": r"(education|academic)",
    "skills": r"(skills|technical skills|technologies)",
    "projects": r"(projects|project experience)"
}


def extract_resume_text(path: str, ext: str) -> str:
    ext = ext.lower()
    if ext == ".pdf":
        reader = PdfReader(path)
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if ext in {".docx", ".doc"}:
        document = Document(path)
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
    if ext in {".txt", ".md"}:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    raise HTTPException(400, "Unsupported file type for ATS analysis")


def analyze_resume_text(text: str) -> dict:
    lowered = text.lower()
    words = re.findall(r"\b[\w.+#/-]+\b", text)
    word_count = len(words)
    email_present = bool(re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text))
    phone_present = bool(re.search(r"(\+\d{1,3}[\s-]?)?(\d[\s-]?){10,}", text))
    linkedin_present = "linkedin.com" in lowered
    github_present = "github.com" in lowered
    metrics_count = len(re.findall(r"\b\d+%|\b\d+\+|\$\d+|\b\d+\b", text))
    action_verbs_found = sorted({verb for verb in ACTION_VERBS if re.search(rf"\b{re.escape(verb)}\b", lowered)})
    keywords_found = sorted({kw for kw in COMMON_KEYWORDS if kw in lowered})
    sections_found = [name for name, pattern in SECTION_PATTERNS.items() if re.search(pattern, lowered)]
    all_caps_ratio = (sum(1 for char in text if char.isupper()) / max(1, sum(1 for char in text if char.isalpha())))

    score = 35
    score += min(15, len(sections_found) * 3)
    score += 10 if email_present else 0
    score += 6 if phone_present else 0
    score += 4 if linkedin_present else 0
    score += 3 if github_present else 0
    score += min(12, len(keywords_found) * 2)
    score += min(10, metrics_count * 2)
    score += min(8, len(action_verbs_found))

    if word_count < 200:
        score -= 12
    elif word_count < 350:
        score -= 5
    elif word_count > 1200:
        score -= 8

    if all_caps_ratio > 0.18:
        score -= 5

    score = max(0, min(100, score))

    suggestions = []
    if not email_present or not phone_present:
        suggestions.append("Add email, phone, and LinkedIn at the top.")
    if "projects" not in sections_found:
        suggestions.append("Add a Projects section with tools and outcomes.")
    if metrics_count < 3:
        suggestions.append("Add numbers like %, time saved, revenue, or users served.")
    if len(keywords_found) < 5:
        suggestions.append("Match more job keywords from the target role.")
    if len(action_verbs_found) < 4:
        suggestions.append("Start bullets with strong verbs like built, led, or improved.")
    if word_count < 350:
        suggestions.append("Add more relevant detail to reach stronger ATS coverage.")

    bonus_suggestions = []
    if "summary" not in sections_found:
        bonus_suggestions.append("Add a 2-line summary tailored to the role.")
    if len(keywords_found) >= 5:
        bonus_suggestions.append("Mirror exact skill names from the job description.")
    if metrics_count >= 3:
        bonus_suggestions.append("Move your strongest quantified wins to the top third.")
    if "skills" in sections_found:
        bonus_suggestions.append("Group skills by language, framework, cloud, and tools.")
    if not github_present:
        bonus_suggestions.append("Add GitHub or portfolio links if they support the role.")

    suggestions.extend(item for item in bonus_suggestions if item not in suggestions)

    return {
        "ats_score": score,
        "word_count": word_count,
        "keywords_found": keywords_found,
        "sections_found": sections_found,
        "metrics_count": metrics_count,
        "action_verbs_found": action_verbs_found,
        "suggestions": suggestions[:5],
    }


def clean_resume_lines(text: str) -> list[str]:
    return [line.strip() for line in text.splitlines() if line.strip()]


def extract_contact_profile(text: str) -> dict:
    lines = clean_resume_lines(text)
    email_match = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)
    phone_match = re.search(r"(\+\d{1,3}[\s-]?)?(\d[\s-]?){10,}", text)
    linkedin_match = re.search(r"(?:https?://)?(?:www\.)?linkedin\.com/[^\s|,;]+", text, re.I)
    github_match = re.search(r"(?:https?://)?(?:www\.)?github\.com/[^\s|,;]+", text, re.I)
    name = lines[0] if lines else ""
    if email_match and name == email_match.group(0):
        name = ""
    return {
        "name": name,
        "phone": phone_match.group(0).strip() if phone_match else "",
        "email": email_match.group(0).strip() if email_match else "",
        "linkedin": linkedin_match.group(0).strip() if linkedin_match else "",
        "github": github_match.group(0).strip() if github_match else "",
    }


def section_text_between(text: str, start_pattern: str, end_patterns: list[str]) -> str:
    match = re.search(start_pattern, text, re.I)
    if not match:
        return ""
    start = match.end()
    end = len(text)
    for pattern in end_patterns:
        next_match = re.search(pattern, text[start:], re.I)
        if next_match:
            end = min(end, start + next_match.start())
    return text[start:end].strip()


def bullets_from_text(text: str, limit: int = 12) -> list[str]:
    lines = clean_resume_lines(text)
    bullets = []
    for line in lines:
        cleaned = re.sub(r"^[•\-\*\u2022]\s*", "", line).strip()
        if not cleaned:
            continue
        if len(cleaned.split()) < 3 and ":" not in cleaned:
            continue
        bullets.append(cleaned)
        if len(bullets) >= limit:
            break
    return bullets


def split_list_field(value: str) -> list[str]:
    return [item.strip() for item in re.split(r"[\n;]+", value or "") if item.strip()]


def extract_reference_resume_data(text: str, overrides: dict | None = None) -> dict:
    overrides = overrides or {}
    profile = extract_contact_profile(text)
    profile.update({key: str(value).strip() for key, value in overrides.items() if key in profile and str(value).strip()})

    education_text = section_text_between(
        text,
        r"\bEducation\b",
        [r"\bCompetitive Programming\b", r"\bProjects\b", r"\bTechnical Skills\b", r"\bExperience\b"],
    )
    competitive_text = section_text_between(
        text,
        r"\b(Competitive Programming|Coding Profiles|Achievements)\b",
        [r"\bProjects\b", r"\bTechnical Skills\b", r"\bExperience\b"],
    )
    projects_text = section_text_between(
        text,
        r"\bProjects?\b",
        [r"\bTechnical Skills\b", r"\bSkills\b", r"\bExperience\b", r"\bEducation\b"],
    )
    skills_text = section_text_between(
        text,
        r"\b(Technical Skills|Skills)\b",
        [r"\bProjects?\b", r"\bExperience\b", r"\bEducation\b", r"\bCertifications?\b"],
    )

    data = {
        **profile,
        "education": bullets_from_text(education_text, 8),
        "competitive_programming": bullets_from_text(competitive_text, 6),
        "projects": bullets_from_text(projects_text, 12),
        "technical_skills": bullets_from_text(skills_text, 10),
    }

    for key in ["education", "competitive_programming", "projects", "technical_skills"]:
        if overrides.get(key):
            data[key] = split_list_field(str(overrides[key]))

    return data


def missing_reference_fields(data: dict) -> list[dict]:
    fields = [
        ("name", "Full name", "Ankit Kumar"),
        ("phone", "Phone", "+91 9876543210"),
        ("email", "Email", "name@example.com"),
        ("linkedin", "LinkedIn", "linkedin.com/in/username"),
        ("github", "GitHub", "github.com/username"),
        ("education", "Education", "College, degree, graduation year; Class XII; Class X"),
        ("projects", "Projects", "Project name - short impact bullets; another project - bullets"),
        ("technical_skills", "Technical skills", "Programming Languages: Python, C++; Tools: Git, Docker"),
    ]
    missing = []
    for key, label, placeholder in fields:
        value = data.get(key)
        if not value:
            missing.append({"key": key, "label": label, "placeholder": placeholder})
    return missing


def add_doc_heading(document: Document, text: str):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(2)


def add_doc_bullet(document: Document, text: str):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(text)
    run.font.size = Pt(9)


def build_reference_docx(data: dict, output_path: str):
    document = Document()
    section = document.sections[0]
    section.top_margin = Pt(28)
    section.bottom_margin = Pt(28)
    section.left_margin = Pt(36)
    section.right_margin = Pt(36)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(9)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(data.get("name", "Candidate Name"))
    title_run.bold = True
    title_run.font.size = Pt(16)

    contact_parts = [data.get("phone"), data.get("email"), data.get("linkedin"), data.get("github")]
    contact = document.add_paragraph(" | ".join(part for part in contact_parts if part))
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(6)

    add_doc_heading(document, "Education")
    for item in data.get("education", []):
        add_doc_bullet(document, item)

    if data.get("competitive_programming"):
        add_doc_heading(document, "Competitive Programming & DSA")
        for item in data["competitive_programming"]:
            add_doc_bullet(document, item)

    add_doc_heading(document, "Projects")
    for item in data.get("projects", []):
        add_doc_bullet(document, item)

    add_doc_heading(document, "Technical Skills")
    for item in data.get("technical_skills", []):
        add_doc_bullet(document, item)

    document.save(output_path)


def reference_resume_suggestions(data: dict) -> list[str]:
    suggestions = [
        "Use a single centered name and contact line like the reference resume.",
        "Keep sections in this order: Education, Competitive Programming & DSA, Projects, Technical Skills.",
        "Convert long paragraphs into short impact bullets.",
        "Group skills by category instead of listing one long comma-separated block.",
    ]
    if not data.get("competitive_programming"):
        suggestions.append("Add coding profiles or DSA achievements if available; otherwise omit that section.")
    if len(data.get("projects", [])) < 3:
        suggestions.append("Add at least three projects with technology and outcome bullets to match the reference depth.")
    return suggestions[:6]


def normalize_skill_terms(raw_skills) -> set[str]:
    if isinstance(raw_skills, str):
        try:
            raw_skills = json.loads(raw_skills)
        except json.JSONDecodeError:
            raw_skills = [item.strip() for item in raw_skills.split(",") if item.strip()]
    if not isinstance(raw_skills, list):
        return set()
    return {str(skill).strip().lower() for skill in raw_skills if str(skill).strip()}


def parse_json_list(raw_value, fallback=None) -> list:
    fallback = fallback or []
    if isinstance(raw_value, list):
        return raw_value
    try:
        value = json.loads(raw_value or "[]")
    except (TypeError, json.JSONDecodeError):
        return fallback
    return value if isinstance(value, list) else fallback


def parse_salary_lpa(text: str) -> int:
    lowered = (text or "").lower()
    lpa_match = re.search(r"(\d{1,3})(?:\s*[-–]\s*(\d{1,3}))?\s*(?:lpa|lakhs?|lakh)", lowered)
    if lpa_match:
        return int(lpa_match.group(2) or lpa_match.group(1))
    inr_match = re.search(r"₹\s*(\d{1,3})(?:\s*[-–]\s*(\d{1,3}))?", lowered)
    if inr_match:
        return int(inr_match.group(2) or inr_match.group(1))
    usd_match = re.search(r"\$?\s*(\d{2,3})\s*k", lowered)
    if usd_match:
        return int(round(int(usd_match.group(1)) * 0.83))
    return 0


def company_package_floor(company_name: str) -> int:
    premium = {name for name, _, _ in TOP_COMPANY_SEEDS[:80]}
    return 18 if company_name in premium else 10


def preference_matches_job(job_row: sqlite3.Row | dict, preferences: dict | None) -> bool:
    if not preferences:
        return True
    selected_companies = set(preferences.get("company_names") or [])
    if selected_companies and job_row["company"] not in selected_companies:
        return False
    min_salary = int(preferences.get("min_salary_lpa") or 0)
    if min_salary:
        offered = parse_salary_lpa(job_row["package_text"] or job_row["salary"] or "")
        if offered and offered < min_salary:
            return False
        if not offered and company_package_floor(job_row["company"]) < min_salary:
            return False
    roles = [role.strip().lower() for role in (preferences.get("roles") or "").split(",") if role.strip()]
    if roles and not any(role in (job_row["title"] or "").lower() for role in roles):
        return False
    locations = [loc.strip().lower() for loc in (preferences.get("locations") or "").split(",") if loc.strip()]
    if locations and not any(loc in (job_row["location"] or "").lower() for loc in locations):
        return False
    return True


def get_user_job_preferences(conn: sqlite3.Connection, user_id: str) -> dict:
    row = conn.execute("SELECT * FROM user_job_preferences WHERE user_id=?", (user_id,)).fetchone()
    if not row:
        return {"company_ids": [], "company_names": [], "min_salary_lpa": 0, "roles": "", "locations": ""}
    company_ids = [int(item) for item in parse_json_list(row["company_ids"]) if str(item).isdigit()]
    names = []
    if company_ids:
        placeholders = ",".join("?" for _ in company_ids)
        names = [
            item["name"] for item in conn.execute(
                f"SELECT name FROM company_watchlist WHERE id IN ({placeholders})",
                company_ids,
            ).fetchall()
        ]
    return {
        "company_ids": company_ids,
        "company_names": names,
        "min_salary_lpa": int(row["min_salary_lpa"] or 0),
        "roles": row["roles"] or "",
        "locations": row["locations"] or "",
    }


class CareerLinkParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.links = []
        self._active_href = ""
        self._active_text = []

    def handle_starttag(self, tag, attrs):
        if tag != "a":
            return
        attrs_map = dict(attrs)
        self._active_href = attrs_map.get("href", "")
        self._active_text = []

    def handle_data(self, data):
        if self._active_href:
            self._active_text.append(data.strip())

    def handle_endtag(self, tag):
        if tag == "a" and self._active_href:
            text = " ".join(part for part in self._active_text if part).strip()
            self.links.append({"href": self._active_href, "text": text})
            self._active_href = ""
            self._active_text = []


def fetch_career_links(source: sqlite3.Row, limit: int = 8) -> list[dict]:
    req = urlrequest.Request(
        source["careers_url"],
        headers={"User-Agent": "FitScoreJobBot/1.0 (+https://github.com/ankit-kumar-developer-122/FitScore)"},
    )
    try:
        with urlrequest.urlopen(req, timeout=8) as response:
            html = response.read(750000).decode("utf-8", errors="ignore")
    except Exception:
        return []

    parser = CareerLinkParser()
    parser.feed(html)
    found = []
    seen = set()
    for link in parser.links:
        text = re.sub(r"\s+", " ", link["text"] or "").strip()
        href = urljoin(source["careers_url"], link["href"])
        haystack = f"{text} {href}".lower()
        if not text or href in seen:
            continue
        if not any(keyword in haystack for keyword in JOB_TITLE_KEYWORDS):
            continue
        seen.add(href)
        found.append({
            "title": text[:120],
            "company": source["name"],
            "application_url": href,
            "source_name": source["name"],
            "source_url": source["careers_url"],
        })
        if len(found) >= limit:
            break
    return found


def infer_skills_from_title(title: str) -> list[str]:
    lowered = title.lower()
    skills = []
    mapping = {
        "frontend": ["JavaScript", "React", "CSS"],
        "backend": ["Python", "SQL", "API"],
        "full": ["JavaScript", "Python", "SQL"],
        "data": ["Python", "SQL", "Data"],
        "machine learning": ["Python", "ML", "Data"],
        "ml": ["Python", "ML"],
        "devops": ["AWS", "Docker", "Kubernetes"],
        "cloud": ["AWS", "Docker"],
        "security": ["Security", "Cloud"],
        "android": ["Kotlin", "Android"],
        "ios": ["Swift", "iOS"],
    }
    for key, values in mapping.items():
        if key in lowered:
            skills.extend(values)
    return sorted(set(skills or ["Python", "JavaScript", "SQL"]))


def upsert_scraped_jobs(conn: sqlite3.Connection, jobs: list[dict]) -> dict:
    inserted = 0
    skipped = 0
    inserted_job_ids = []
    now = datetime.utcnow().isoformat()
    for job in jobs:
        prior = conn.execute("SELECT id FROM jobs WHERE application_url=?", (job["application_url"],)).fetchone()
        if prior:
            skipped += 1
            continue
        company = job["company"]
        floor = company_package_floor(company)
        package_text = f"{floor}+ LPA expected"
        cursor = conn.execute(
            """
            INSERT INTO jobs (title,company,salary,match_pct,status,skills,experience,description,location,package_text,application_url,source_name,source_url,created_at)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """,
            (
                job["title"],
                company,
                package_text,
                80,
                "active",
                json.dumps(infer_skills_from_title(job["title"])),
                "See posting",
                f"Imported from {job['source_name']} careers/jobs source.",
                "Remote / India / Global",
                package_text,
                job["application_url"],
                job["source_name"],
                job["source_url"],
                now,
            )
        )
        inserted += 1
        inserted_job_ids.append(cursor.lastrowid)
    return {"inserted": inserted, "skipped": skipped, "inserted_job_ids": inserted_job_ids}


def get_user_match_profile(conn: sqlite3.Connection, user_row: sqlite3.Row) -> dict:
    stored_skills = normalize_skill_terms(user_row["skills"] or "[]")
    latest_resume = conn.execute(
        "SELECT ats_score, analysis_json FROM resumes WHERE user_id=? ORDER BY uploaded_at DESC LIMIT 1",
        (user_row["id"],)
    ).fetchone()
    resume_analysis = {}
    if latest_resume and latest_resume["analysis_json"]:
        try:
            resume_analysis = json.loads(latest_resume["analysis_json"])
        except json.JSONDecodeError:
            resume_analysis = {}
    resume_keywords = {str(item).strip().lower() for item in resume_analysis.get("keywords_found", []) if str(item).strip()}
    combined_skills = sorted(stored_skills | resume_keywords)
    return {
        "skills": combined_skills,
        "ats_score": int((latest_resume["ats_score"] if latest_resume and latest_resume["ats_score"] is not None else user_row["ats_score"]) or 0),
    }


def calculate_job_match(job_skills: set[str], user_profile: dict) -> int:
    user_skills = set(user_profile["skills"])
    if not user_skills:
        return 0
    overlap = len(job_skills & user_skills)
    skill_score = int((overlap / max(1, len(job_skills))) * 70) if job_skills else 0
    ats_bonus = min(30, int(user_profile["ats_score"] * 0.3))
    return min(100, skill_score + ats_bonus)


def build_match_payload(user_row: sqlite3.Row, job_row: sqlite3.Row, match_score: int, user_profile: dict) -> dict:
    return {
        "user": {
            "id": user_row["id"],
            "name": user_row["name"],
            "email": user_row["email"],
            "role": user_row["role"],
            "ats_score": user_profile["ats_score"],
        },
        "job": {
            "id": job_row["id"],
            "title": job_row["title"],
            "company": job_row["company"],
            "salary": job_row["salary"],
            "location": job_row["location"],
            "package_text": job_row["package_text"],
            "experience": job_row["experience"],
            "description": job_row["description"],
            "application_url": job_row["application_url"],
        },
        "match": {
            "score": match_score,
            "shared_skills": sorted(set(user_profile["skills"]) & normalize_skill_terms(job_row["skills"] or "[]")),
        },
        "sent_at": datetime.utcnow().isoformat(),
    }


def send_n8n_webhook(payload: dict) -> tuple[bool, str]:
    if not N8N_WEBHOOK_URL:
        return False, "N8N webhook URL is not configured"
    req = urlrequest.Request(
        N8N_WEBHOOK_URL,
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Content-Type": "application/json",
            **({"X-FitScore-Secret": N8N_WEBHOOK_SECRET} if N8N_WEBHOOK_SECRET else {})
        },
        method="POST",
    )
    try:
        with urlrequest.urlopen(req, timeout=10) as response:
            return True, f"Webhook accepted with status {response.status}"
    except HTTPError as exc:
        return False, f"Webhook HTTP error {exc.code}"
    except URLError as exc:
        return False, f"Webhook connection error: {exc.reason}"


def notify_matches_for_job(conn: sqlite3.Connection, job_id: int, auto_email_only: bool = True) -> dict:
    job_row = conn.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()
    if not job_row:
        raise HTTPException(404, "Job not found")

    user_query = "SELECT * FROM users WHERE role='candidate'"
    if auto_email_only:
        user_query += " AND auto_email=1"
    user_rows = conn.execute(user_query).fetchall()

    job_skills = normalize_skill_terms(job_row["skills"] or "[]")
    notified = []
    skipped = 0

    for user_row in user_rows:
        preferences = get_user_job_preferences(conn, user_row["id"])
        if not preference_matches_job(job_row, preferences):
            skipped += 1
            continue

        prior = conn.execute(
            "SELECT 1 FROM notification_logs WHERE user_id=? AND job_id=? AND channel='n8n_email' AND status='sent'",
            (user_row["id"], job_id)
        ).fetchone()
        if prior:
            skipped += 1
            continue

        user_profile = get_user_match_profile(conn, user_row)
        match_score = calculate_job_match(job_skills, user_profile)
        if match_score < MATCH_THRESHOLD:
            continue

        payload = build_match_payload(user_row, job_row, match_score, user_profile)
        ok, message = send_n8n_webhook(payload)
        conn.execute(
            """
            INSERT INTO notification_logs (user_id, job_id, channel, status, payload_json, created_at)
            VALUES (?,?,?,?,?,?)
            ON CONFLICT(user_id, job_id, channel) DO UPDATE SET
                status=excluded.status,
                payload_json=excluded.payload_json,
                created_at=excluded.created_at
            """,
            (
                user_row["id"],
                job_id,
                "n8n_email",
                "sent" if ok else "failed",
                json.dumps({"payload": payload, "message": message}),
                datetime.utcnow().isoformat(),
            )
        )
        if ok:
            notified.append({
                "user_id": user_row["id"],
                "email": user_row["email"],
                "match_score": match_score,
            })

    return {
        "job_id": job_id,
        "job_title": job_row["title"],
        "notified_count": len(notified),
        "skipped_count": skipped,
        "notified_users": notified,
    }

# ── Auth Endpoints ───────────────────────────────────────────────────
def hash_pw(pw):
    return hashlib.sha256(pw.encode()).hexdigest()

@app.post("/api/register")
async def register(name: str = Form(...), email: str = Form(...), password: str = Form(...), role: str = Form("candidate")):
    conn = get_db()
    uid = str(uuid.uuid4())
    try:
        conn.execute("INSERT INTO users (id,name,email,password_hash,role,fitscore,ats_score,skills,created_at) VALUES (?,?,?,?,?,?,?,?,?)",
                     (uid, name, email, hash_pw(password), role, 85, 92, '["HTML","CSS","Python","Data Analysis"]', datetime.utcnow().isoformat()))
        conn.commit()
    except sqlite3.IntegrityError:
        conn.close()
        raise HTTPException(400, "Email already registered")
    conn.close()
    return {"id": uid, "name": name, "email": email, "role": role}

@app.post("/api/login")
async def login(email: str = Form(...), password: str = Form(...)):
    conn = get_db()
    user = conn.execute("SELECT * FROM users WHERE email=? AND password_hash=?", (email, hash_pw(password))).fetchone()
    conn.close()
    if not user:
        raise HTTPException(401, "Invalid credentials")
    return dict(user)

# ── User Endpoints ───────────────────────────────────────────────────
@app.get("/api/users")
async def get_users():
    conn = get_db()
    rows = conn.execute("SELECT id,name,email,role,fitscore,ats_score,created_at FROM users").fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.post("/api/users/{user_id}/toggle-email")
async def toggle_email(user_id: str):
    conn = get_db()
    user = conn.execute("SELECT auto_email FROM users WHERE id=?", (user_id,)).fetchone()
    if not user:
        conn.close()
        raise HTTPException(404, "User not found")
    new_val = 0 if user["auto_email"] else 1
    conn.execute("UPDATE users SET auto_email=? WHERE id=?", (new_val, user_id))
    conn.commit()
    conn.close()
    return {"auto_email": bool(new_val)}


@app.get("/api/companies")
async def get_companies():
    conn = get_db()
    rows = conn.execute(
        "SELECT * FROM company_watchlist ORDER BY source_type, package_tier DESC, name"
    ).fetchall()
    conn.close()
    return [dict(row) for row in rows]


@app.get("/api/users/{user_id}/job-preferences")
async def read_job_preferences(user_id: str):
    conn = get_db()
    preferences = get_user_job_preferences(conn, user_id)
    conn.close()
    return preferences


@app.post("/api/users/{user_id}/job-preferences")
async def save_job_preferences(
    user_id: str,
    company_ids: str = Form("[]"),
    min_salary_lpa: int = Form(0),
    roles: str = Form(""),
    locations: str = Form(""),
):
    ids = [int(item) for item in parse_json_list(company_ids) if str(item).isdigit()]
    conn = get_db()
    if ids:
        placeholders = ",".join("?" for _ in ids)
        valid_ids = {
            row["id"] for row in conn.execute(
                f"SELECT id FROM company_watchlist WHERE id IN ({placeholders})",
                ids,
            ).fetchall()
        }
        ids = [item for item in ids if item in valid_ids]
    conn.execute(
        """
        INSERT INTO user_job_preferences (user_id,company_ids,min_salary_lpa,roles,locations,updated_at)
        VALUES (?,?,?,?,?,?)
        ON CONFLICT(user_id) DO UPDATE SET
            company_ids=excluded.company_ids,
            min_salary_lpa=excluded.min_salary_lpa,
            roles=excluded.roles,
            locations=excluded.locations,
            updated_at=excluded.updated_at
        """,
        (
            user_id,
            json.dumps(ids),
            max(0, int(min_salary_lpa or 0)),
            roles.strip(),
            locations.strip(),
            datetime.utcnow().isoformat(),
        ),
    )
    conn.commit()
    preferences = get_user_job_preferences(conn, user_id)
    conn.close()
    return preferences

# ── Job Endpoints ────────────────────────────────────────────────────
@app.get("/api/jobs")
async def get_jobs(user_id: str = ""):
    conn = get_db()
    preferences = get_user_job_preferences(conn, user_id) if user_id else None
    rows = conn.execute("SELECT * FROM jobs ORDER BY created_at DESC, match_pct DESC").fetchall()
    conn.close()
    return [dict(row) for row in rows if preference_matches_job(row, preferences)]

@app.post("/api/jobs")
async def create_job(request: Request, title: str = Form(...), company: str = Form(...), salary: str = Form(""), skills: str = Form("[]"), experience: str = Form(""), description: str = Form(""), location: str = Form(""), package_text: str = Form(""), application_url: str = Form("")):
    require_roles(request, {"recruiter"})
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO jobs (title,company,salary,skills,experience,description,location,package_text,application_url,created_at) VALUES (?,?,?,?,?,?,?,?,?,?)",
        (title, company, salary, skills, experience, description, location, package_text, application_url, datetime.utcnow().isoformat())
    )
    job_id = cursor.lastrowid
    notification_result = notify_matches_for_job(conn, job_id)
    conn.commit()
    conn.close()
    return {"ok": True, "job_id": job_id, "notifications": notification_result}


@app.post("/api/jobs/refresh-careers")
async def refresh_career_jobs(request: Request, company_ids: str = Form("[]"), max_sources: int = Form(25)):
    user = get_request_user(request)
    ids = [int(item) for item in parse_json_list(company_ids) if str(item).isdigit()]
    conn = get_db()
    if not ids and user["role"] == "candidate":
        ids = get_user_job_preferences(conn, user["id"]).get("company_ids", [])
    if ids:
        placeholders = ",".join("?" for _ in ids[:max_sources])
        sources = conn.execute(
            f"SELECT * FROM company_watchlist WHERE id IN ({placeholders})",
            ids[:max_sources],
        ).fetchall()
    else:
        sources = conn.execute(
            "SELECT * FROM company_watchlist ORDER BY source_type, package_tier DESC, name LIMIT ?",
            (max(1, min(int(max_sources or 25), 50)),)
        ).fetchall()

    scraped = []
    for source in sources:
        scraped.extend(fetch_career_links(source, limit=5))

    result = upsert_scraped_jobs(conn, scraped)
    notification_results = [
        notify_matches_for_job(conn, job_id)
        for job_id in result.get("inserted_job_ids", [])
    ]
    conn.commit()
    conn.close()
    return {
        "inserted": result["inserted"],
        "skipped": result["skipped"],
        "sources_checked": len(sources),
        "jobs_found": len(scraped),
        "notifications_sent": sum(item.get("notified_count", 0) for item in notification_results),
    }

# ── Resume Upload ────────────────────────────────────────────────────
@app.post("/api/resumes/upload")
async def upload_resume(user_id: str = Form(...), file: UploadFile = File(...), version_tag: str = Form("Original")):
    ext = os.path.splitext(file.filename)[1]
    safe_name = f"{uuid.uuid4().hex}{ext}"
    path = os.path.join("uploads", safe_name)
    contents = await file.read()
    with open(path, "wb") as f:
        f.write(contents)
    extracted_text = extract_resume_text(path, ext)
    analysis = analyze_resume_text(extracted_text)
    conn = get_db()
    conn.execute("INSERT INTO resumes (user_id,filename,filepath,version_tag,uploaded_at) VALUES (?,?,?,?,?)",
                 (user_id, file.filename, path, version_tag, datetime.utcnow().isoformat()))
    conn.execute(
        "UPDATE resumes SET ats_score=?, analysis_json=? WHERE id = last_insert_rowid()",
        (analysis["ats_score"], json.dumps(analysis))
    )
    conn.execute("UPDATE users SET ats_score=? WHERE id=?", (analysis["ats_score"], user_id))
    conn.commit()
    conn.close()
    return {
        "filename": file.filename,
        "path": path,
        "tag": version_tag,
        "ats_score": analysis["ats_score"],
        "analysis": analysis,
    }

@app.get("/api/resumes/{user_id}")
async def get_resumes(user_id: str):
    conn = get_db()
    rows = conn.execute("SELECT * FROM resumes WHERE user_id=? ORDER BY uploaded_at DESC", (user_id,)).fetchall()
    conn.close()
    resumes = []
    for row in rows:
        item = dict(row)
        try:
            item["analysis"] = json.loads(item.get("analysis_json") or "{}")
        except json.JSONDecodeError:
            item["analysis"] = {}
        resumes.append(item)
    return resumes


@app.post("/api/resumes/{resume_id}/make-like-reference")
async def make_resume_like_reference(resume_id: int, missing_json: str = Form("{}")):
    conn = get_db()
    resume = conn.execute("SELECT * FROM resumes WHERE id=?", (resume_id,)).fetchone()
    conn.close()
    if not resume:
        raise HTTPException(404, "Resume not found")

    filepath = resume["filepath"]
    if not filepath or not os.path.exists(filepath):
        raise HTTPException(404, "Uploaded resume file is missing")

    try:
        overrides = json.loads(missing_json or "{}")
    except json.JSONDecodeError:
        raise HTTPException(400, "Missing field data must be valid JSON")

    ext = os.path.splitext(filepath)[1]
    text = extract_resume_text(filepath, ext)
    data = extract_reference_resume_data(text, overrides)
    missing = missing_reference_fields(data)
    suggestions = reference_resume_suggestions(data)

    if missing:
        return {
            "requires_input": True,
            "missing_fields": missing,
            "suggestions": suggestions,
        }

    output_name = f"reference_resume_{resume_id}_{uuid.uuid4().hex[:8]}.docx"
    output_path = os.path.join("uploads", "generated", output_name)
    build_reference_docx(data, output_path)

    analysis = analyze_resume_text(extract_resume_text(output_path, ".docx"))
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO resumes (user_id,filename,filepath,version_tag,ats_score,analysis_json,uploaded_at) VALUES (?,?,?,?,?,?,?)",
        (
            resume["user_id"],
            output_name,
            output_path,
            "Reference Style",
            analysis["ats_score"],
            json.dumps(analysis),
            datetime.utcnow().isoformat(),
        )
    )
    generated_resume_id = cursor.lastrowid
    conn.commit()
    conn.close()

    return {
        "requires_input": False,
        "download_url": f"/api/resume-files/{generated_resume_id}/download",
        "filename": output_name,
        "resume_id": generated_resume_id,
        "suggestions": suggestions,
    }


@app.get("/api/resume-files/{resume_id}/download")
async def download_resume_file(resume_id: int):
    conn = get_db()
    resume = conn.execute("SELECT filename, filepath FROM resumes WHERE id=?", (resume_id,)).fetchone()
    conn.close()
    if not resume or not resume["filepath"] or not os.path.exists(resume["filepath"]):
        raise HTTPException(404, "Resume file not found")
    return FileResponse(
        resume["filepath"],
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=resume["filename"],
    )

# ── Applications ─────────────────────────────────────────────────────
@app.get("/api/applications")
async def get_applications():
    conn = get_db()
    rows = conn.execute("SELECT a.*, u.name as user_name, j.title as job_title FROM applications a JOIN users u ON a.user_id=u.id JOIN jobs j ON a.job_id=j.id").fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.post("/api/applications")
async def apply_job(user_id: str = Form(...), job_id: int = Form(...)):
    conn = get_db()
    conn.execute("INSERT INTO applications (user_id,job_id,stage,created_at) VALUES (?,?,?,?)",
                 (user_id, job_id, "applied", datetime.utcnow().isoformat()))
    conn.commit()
    conn.close()
    return {"ok": True}


@app.post("/api/automation/match-and-notify")
async def match_and_notify(request: Request, job_id: int = Form(...), auto_email_only: int = Form(1)):
    require_roles(request, {"recruiter"})
    conn = get_db()
    result = notify_matches_for_job(conn, job_id, auto_email_only=bool(auto_email_only))
    conn.commit()
    conn.close()
    return result

# ── Workflows ────────────────────────────────────────────────────────
@app.get("/api/workflows")
async def get_workflows(request: Request):
    require_roles(request, {"recruiter"})
    conn = get_db()
    rows = conn.execute("SELECT * FROM workflows").fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.get("/api/recruiter/dashboard")
async def recruiter_dashboard(request: Request, job_id: int | None = None):
    require_roles(request, {"recruiter"})
    conn = get_db()
    jobs = [dict(row) for row in conn.execute("SELECT * FROM jobs ORDER BY created_at DESC, id DESC").fetchall()]
    if not jobs:
        conn.close()
        return {
            "jobs": [],
            "selected_job": None,
            "candidates": [],
            "summary": {"total_jobs": 0, "total_candidates": 0, "shortlisted": 0, "avg_ats": 0},
        }

    selected_job = next((job for job in jobs if job["id"] == job_id), jobs[0] if jobs else None)
    job_skills = normalize_skill_terms(selected_job["skills"] or "[]")
    candidate_rows = conn.execute("SELECT * FROM users WHERE role='candidate'").fetchall()

    ranked_candidates = []
    for candidate in candidate_rows:
        profile = get_user_match_profile(conn, candidate)
        match_score = calculate_job_match(job_skills, profile)
        potential_score = round((match_score * 0.7) + (profile["ats_score"] * 0.3), 1)
        ranked_candidates.append({
            "id": candidate["id"],
            "name": candidate["name"],
            "email": candidate["email"],
            "ats_score": profile["ats_score"],
            "match_score": match_score,
            "potential_score": potential_score,
            "skills": profile["skills"][:8],
            "shared_skills": sorted(set(profile["skills"]) & job_skills),
        })

    ranked_candidates.sort(key=lambda item: (-item["ats_score"], -item["match_score"], item["name"].lower()))
    top_candidates = ranked_candidates[:10]
    high_potential_low_ats = sorted(
        [
            candidate for candidate in ranked_candidates
            if candidate["match_score"] >= 60 and candidate["ats_score"] < 75
        ],
        key=lambda item: (-item["match_score"], -item["potential_score"], item["name"].lower())
    )[:10]
    shortlisted = [candidate for candidate in ranked_candidates if candidate["ats_score"] >= 80]
    avg_ats = round(sum(candidate["ats_score"] for candidate in ranked_candidates) / len(ranked_candidates), 1) if ranked_candidates else 0

    conn.close()
    return {
        "jobs": jobs,
        "selected_job": selected_job,
        "candidates": ranked_candidates,
        "top_candidates": top_candidates,
        "high_potential_low_ats": high_potential_low_ats,
        "summary": {
            "total_jobs": len(jobs),
            "total_candidates": len(ranked_candidates),
            "shortlisted": len(shortlisted),
            "avg_ats": avg_ats,
        },
    }

# ── Analytics ────────────────────────────────────────────────────────
@app.get("/api/analytics/summary")
async def analytics_summary():
    conn = get_db()
    total_users = conn.execute("SELECT COUNT(*) FROM users").fetchone()[0]
    total_jobs = conn.execute("SELECT COUNT(*) FROM jobs").fetchone()[0]
    total_apps = conn.execute("SELECT COUNT(*) FROM applications").fetchone()[0]
    total_resumes = conn.execute("SELECT COUNT(*) FROM resumes").fetchone()[0]
    conn.close()
    return {"total_users": total_users, "total_jobs": total_jobs, "total_applications": total_apps, "total_resumes": total_resumes}


@app.get("/api/health")
async def health_check():
    return {"status": "ok"}

# ── DB Tables (admin view) ──────────────────────────────────────────
@app.get("/api/db/tables")
async def db_tables(request: Request):
    require_roles(request, {"recruiter"})
    conn = get_db()
    tables = conn.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()
    result = {}
    for t in tables:
        name = t["name"]
        rows = conn.execute(f"SELECT * FROM {name} LIMIT 50").fetchall()
        result[name] = [dict(r) for r in rows]
    conn.close()
    return result

# ── Serve login page ─────────────────────────────────────────────────
@app.get("/login")
async def login_page():
    return FileResponse("public/login.html", headers={"Cache-Control": "no-cache"})

# ── Static Files (MUST be last) ─────────────────────────────────────
app.mount("/", CachedStaticFiles(directory="public", html=True), name="public")
